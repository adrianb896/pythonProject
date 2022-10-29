import docx
import re
import xlwings
import pandas as pd
from tabulate import tabulate


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    fullText = [ele for ele in fullText if ele.strip()]
    return fullText


hdsLst = getText('/Users/adrian/Desktop/SampledocsTandem/HDS_new_pump.docx')
hrsLst = getText('/Users/adrian/Desktop/SampledocsTandem/HRS_new_pump.docx')
htpLst = getText('/Users/adrian/Desktop/SampledocsTandem/HTP_new_pump.docx')
htrLst = getText('/Users/adrian/Desktop/SampledocsTandem/HTR_new_pump.docx')
prsLst = getText('/Users/adrian/Desktop/SampledocsTandem/PRS_new_pump.docx')
riskLst = getText('/Users/adrian/Desktop/SampledocsTandem/RiskAnalysis_Pump.docx')
sdsLst = getText('/Users/adrian/Desktop/SampledocsTandem/SDS_New_pump_x04.docx')
srsAceLst = getText('/Users/adrian/Desktop/SampledocsTandem/SRS_ACE_Pump_X01.docx')
srsBolusLst = getText('/Users/adrian/Desktop/SampledocsTandem/SRS_BolusCalc_Pump_X04.docx')
srsDoseLst = getText('/Users/adrian/Desktop/SampledocsTandem/SRS_DosingAlgorithm_X03.docx')
svapLst = getText('/Users/adrian/Desktop/SampledocsTandem/SVaP_new_pump.docx')
svatrLst = getText('/Users/adrian/Desktop/SampledocsTandem/SVaTR_new_pump.docx')
svetrLst = getText('/Users/adrian/Desktop/SampledocsTandem/SVeTR_new_pump.docx')
ursLst = getText('/Users/adrian/Desktop/SampledocsTandem/URS_new_pump.docx')

index = 0
ind = []

childrenParent = {'.*:HRD:':('\S*:HRD:\S*', '\S*:HRS:\S*'), '.*:HRS:':('\S*:HRS:\S*', '\S*:PRS:\S*'),
                  '.*:HTP:':('\S*:HTP:\S*', ('\S*:HRS:\S*', '\S*:HRD:\S*')), '.*:HTR:':('\S*:HTR:\S*', '\S*:HTP:\S*'),
                  '.*:PRS:':('\S*:PRS:\S*', ('\S*:URS:\S*', '\S*:RISK:\S*')), '.*:RISK:':('\S*:RISK:\S*'),
                  '.*:SDS:':('\S*:SDS:\S*', '\S*:SRS:\S*'), '.*ACE:SRS:':('\S*:SRS:\S*', ('\S*:PRS:\S*', '\S*:DER:\S*')),
                  '.*BOLUS:SRS:':('\S*:SRS:\S*', '\S*:PRS:\S*'), '.*AID:SRS:':('\S*:SRS:\S*', '\S*:DER:\S*'),
                  '.*:SVAL:':('\S*:SVAL:\S*', '\S*:SRS:\S*'), '.*:SVATR:':('\S*:SVATR:\S*', '\S*:SVAL:\S*'),
                  '.*:UT:':('\S*:UT:\S*', '\S*:UNIT:\S*'), '.*:INS:':('\S*:INS:\S*', '\S*:UNIT:\S*'),
                  '.*:URS:':('\S*:URS:\S*')}


def parseDocument(list):  # this function is able to pass a parent tag with one child tag
    index = 0
    ind = []
    for para in list:  # for each paragraph in the documents look for these list of tags
        for parent, children in childrenParent.items():  # look for items within our dictionary childrenParent
            if re.search(parent, para):  # search for the initial parent tag within each paragraph(line) of a document
                ind.append(index)
                tt = para
                y = re.findall(children[0], para)
                z = re.findall(children[0], para)
                tt = tt.replace(y[0], '')
                tt = tt.replace(z[0], '')
                tt = tt.strip()

                # print(y)
                # print(tt)
                # print(z)
                dict1 = {'Parent Tag': [y],
                         'Info': [tt],
                         'Child Tag': [z]
                          }
                df = pd.DataFrame(dict1)
                print(tabulate(df, headers='keys', tablefmt='fancy_grid'))
            index = index + 1
            # print(ind)


def parseDocument_2(list):  # this function is able to pass a parent tag with two child tags
    index = 0
    ind = []
    for para in list:  # for each paragraph in the documents look for these list of tags
        for parent, children in childrenParent.items():  # look for items within our dictionary childrenParent
            if re.search(parent, para):  # search for the initial child tag within each paragraph(line) of a document
                ind.append(index)
                tt = para
                y = re.findall(children[0], para)
                if y[0].startswith("[") and y[0].endswith("]"):
                    continue
                if re.search(children[1][0], para):
                    z = re.findall(children[1][0], para)
                    tt = tt.replace(y[0], '')
                    tt = tt.replace(z[0], '')
                    tt = tt.strip()
                elif re.search(children[1][1], para):
                    z = re.findall(children[1][1], para)
                    tt = tt.replace(y[0], '')
                    tt = tt.replace(z[0], '')
                    tt = tt.strip()
                dict1 = {'Parent Tag': [y],
                         'Info': [tt],
                         'Child Tag': [z]
                         }
                df = pd.DataFrame(dict1)
                print(tabulate(df, headers='keys', tablefmt='fancy_grid'))
            index = index + 1
            #print(ind)


def parseDocument_3(list):  # this function is able to pass a parent tag with no child tags
    index = 0
    ind = []
    for para in list:  # for each paragraph in the documents look for these list of tags
        for parent, children in childrenParent.items():  # look for items within our dictionary childrenParent
            if re.search(parent, para):  # search for the initial child tag within each paragraph(line) of a document
                ind.append(index)
                tt = para
                y = re.findall(parent, para)
                tt = tt.replace(y[0], '')
                tt = tt.strip()

                dict1 = {'Parent Tag': [y],
                         'Info': [tt]
                         }
                df = pd.DataFrame(dict1)
                print(tabulate(df, headers='keys', tablefmt='fancy_grid'))
            index = index + 1
            # print(ind)

#parseDocument(hdsLst) # works with z having parent[0]
#parseDocument(hrsLst) # works with z having parent[0]
#parseDocument_2(htpLst) # works with second function
#parseDocument(htrLst) # works with z having parent[0]
#parseDocument_2(prsLst) # still needs to be tested and fixed
#parseDocument_3(riskLst) # works with the third function
#parseDocument(sdsLst) # works with z having parent[0]
#parseDocument_2(srsAceLst) # works with the second function
#parseDocument(srsBolusLst) # works with z having parent[0]
#parseDocument_2(srsDoseLst) # still needs to be tested and fixed
#parseDocument(svapLst) # works with z having parent[0]
#parseDocument(svatrLst) # works with z having parent[0]
#parseDocument(svetrLst) # works with z having parent[0]
#parseDocument_3(ursLst) # works with the third function
